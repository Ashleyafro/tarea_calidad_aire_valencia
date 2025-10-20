import requests 
import pandas as pd
from sqlalchemy import create_engine 
import os
import argparse
import matplotlib.pyplot as plt
from pathlib import Path
from docx import Document

    
def check_updates(engine, cleaned_df):

    latest_db_fecha = pd.read_sql("SELECT MAX(fecha_carg) FROM tabla_bd", engine)
    latest_fecha = latest_db_fecha.iloc[0,0]

    # Obtener la fecha max de los datos recolectados de la API
    latest_api_fecha = cleaned_df["fecha_carg"].max()

    # Check si los datos que tenemos están actualizados
    if latest_api_fecha != latest_fecha:
        print("Nuevos datos guardados")
        # Conseguir df más actualizado
        registro_path = "./output/historico/registro_historico.csv"
        #subimos nuestros nuevos datos
        cleaned_df.to_sql("tabla_bd", con = engine, if_exists = 'append', index = False)

        if (os.path.exists(registro_path) and os.path.isfile(registro_path)): #comprueba si la ruta existe y si el archivo que se encuentra en la ruta existe
            os.remove(registro_path)
            registro_df = pd.read_sql("SELECT * FROM tabla_bd", engine) #cargamos nuestra base de datos y la pasamos a csv
            registro_df.to_csv(registro_path, index=False)
            
        else:
            print("No se ha encontrado la ruta del archivo")
            
    else:
        print("No hay nuevos records")
        return None

def save_raw_csv(data, engine):
    # conseguir datos
    df = pd.DataFrame(data["results"])
    cleaned_df = pd.DataFrame(df[["objectid", "nombre", "direccion", "tipozona", "no2", "pm10", "pm25", "tipoemisio", "fecha_carg", "calidad_am", "fiwareid"]])
    cleaned_df = cleaned_df[cleaned_df.objectid != 22]
    #Eliminar filas con valores nulos
    cleaned_df = cleaned_df.dropna(subset=["objectid", "nombre", "direccion", "tipozona", "no2", "pm10", "pm25", "tipoemisio", "fecha_carg", "calidad_am", "fiwareid"])
    cleaned_df["fecha_carg"] = pd.to_datetime(cleaned_df["fecha_carg"], utc = True)

    # Save in PostgreSQL
    file = "./output/historico/registro_historico.csv"
    df_prueba = pd.read_sql("SELECT * FROM tabla_bd", engine)
    if len(df_prueba) == 0: #Si no tenemos BD, creamos una
        cleaned_df.to_sql("tabla_bd", engine, if_exists="replace", index=False)
        print("Datos cargados a PostgreSQL")
        #Condición si la ruta existe pero el archivo no
        if (os.path.exists(file) and not os.path.isfile(file)): #Como esta es nuestra primera BD, guardamelo directamente en el historico
            registro_df = pd.read_sql("SELECT * FROM tabla_bd", engine)
            registro_df.to_csv(file, index=False)
    else:
        check_updates(engine, cleaned_df)
    
    # Guardar el CSV con timestamp incluido en el nombre
    os.makedirs("./data/raw", exist_ok=True)
    latest_date = cleaned_df["fecha_carg"].max() 
    safe_timestamp = latest_date.strftime("%Y-%m-%dT%H-%M-%S")
    csv_filepath = os.path.join("./data/raw", f"ultimo_{safe_timestamp}.csv")

    try:
            cleaned_df.to_csv(csv_filepath, index=False)
            print(f"CSV guardado en: {csv_filepath}")
    except:
            print(f"Error al guardar el archivo CSV")
    
   

def parse_args():
    #Creamos el parser (p) para parsear argumentos 
    p = argparse.ArgumentParser(description="Reportes")
    #Añadimos los argumemtos:
    p.add_argument("--modo", choices=["actual", "historico"], required=True, help="Modo: 'actual' o 'historico'")   #Argumento modo en donde tienes que elegir entre actual o historico
    p.add_argument("--since", type=str, help="Fecha de inicio (formato: YYYY-MM-DD:00:00+0000)") #Argumento que recibe una cadena string de fecha
    p.add_argument("--estacion", type=str, help="ID de estacion")   #Argumento para filtrar segun el fiwareid
    return p.parse_args()

def generate_actual_report():
    # Obtener el último CSV
    csv_folder = "./data/raw/"
    last_file = os.listdir(csv_folder)[0]    #Obtiene el primer nombre del archivo que se encuentra dentro de la carpeta (ruta)
    if not last_file:   #Caso en donde no hayan archivos en esa ruta
        print("No hay CSV")  
    else:
        latest_csv = os.path.join(csv_folder, last_file)
        df_latest = pd.read_csv(latest_csv)     #Pasamos el último csv como un df llamado df_latest
        #Donde exportaremos nuestras gráficas
        output_dir = os.path.join("./output/actual")
        os.makedirs(output_dir, exist_ok=True)

        #Cálculo de valores
        summary = df_latest.groupby("fiwareid")[["no2","pm10","pm25"]].sum()
        os.makedirs("./output/actual", exist_ok=True)
        print("Tabla de resumen guardada")

        #Creación de gráficas

        summary["no2"].plot(kind="bar", title="Gráfico de barras no2 por estación", color = "red")
        plt.xlabel("Estacion")
        plt.ylabel("Nivel NO2") 
        plt.savefig(os.path.join(output_dir, "no2_por_estacion.png"))

        summary["pm10"].plot(kind="bar", title="Gráfico de barras no2 por estación", color = "red")
        plt.xlabel("Estacion")
        plt.ylabel("Nivel pm10")
        plt.savefig(os.path.join(output_dir, "pm10_por_estacion.png"))

        summary["pm25"].plot(kind="bar", title="Gráfico de barras no2 por estación", color = "red")
        plt.xlabel("Estacion")
        plt.ylabel("Nivel pm25")
        plt.savefig(os.path.join(output_dir, "pm25_por_estacion.png"))

        summary.to_csv(os.path.join(output_dir, "tabla_actual.csv"))
        
        #Guardamos las gráficas en un doc
        doc = Document()
        doc.add_heading('Grafica informe actual estacion/no2', level=1)
        doc.add_paragraph('Informe actual sobre el No2 generado según la estación')
        doc.add_picture("./output/actual/no2_por_estacion.png")
        doc.add_heading('Grafica informe actual estacion/pm10', level=1)
        doc.add_paragraph('Informe actual sobre el pm10 generado según la estación')
        doc.add_picture("./output/actual/pm10_por_estacion.png" )
        doc.add_heading('Grafica informe actual estacion/pm25', level=1)
        doc.add_paragraph('Informe actual sobre el pm25 generado según la estación')
        doc.add_picture("./output/actual/pm25_por_estacion.png")

        save_path = os.path.join(output_dir, "Reporte_Actual.docx")
        doc.save(save_path)

        print("Graficos actuales guardados")

def generate_historico_report(since, estacion):
    historico_df = pd.read_csv("./output/historico/registro_historico.csv")
    #El parámetro utc le dice a pandas que trate las fechas como si estuvieran horario UTC
    historico_df["fecha_carg"] = pd.to_datetime(historico_df["fecha_carg"], utc=True) #evitar errores a la hora de comparar fechas (fecha_carg no es tipo datetime)
    output_dir = "./output/historico"
    #Filtramos en base a lo que el usuario haya puesto
    if since and estacion:
        since_date = pd.to_datetime(since, utc=True) #pasamos since a una variable de tipo datetime
        historico_df_filtradofs = historico_df[historico_df["fecha_carg"] >= since_date]   #filtramos desde la fecha dada (since_date) hasta la última fecha de carga
        historico_df_filtradofs = historico_df_filtradofs[historico_df_filtradofs["fiwareid"] == estacion]
        plt.figure(figsize=(10,5))  #tamaño del gráfico (rectangular)
        #Marker muestra un circulo y linestyle une los circulos en una linea contínua
        plt.plot(historico_df_filtradofs["fecha_carg"], historico_df_filtradofs["no2"], marker='o', linestyle='-')    #dibuja gráfica del tiempo en donde "eje x" es la fecha de carga y el "eje y" el no2 generado.

        #Creamos gráfica
        plt.title(f"NO2 histórico de la estación {estacion} desde {since} hasta la fecha más reciente")
        plt.xlabel("Fecha")
        plt.ylabel("Nivel NO2")
        plt.tight_layout()

        #La guardamos
        os.makedirs(output_dir, exist_ok=True)
        plt.savefig(os.path.join(output_dir, "NO2_historico_estacion_since.png"))
        plt.close()
        print("Graficos historicos guardados")

#Crear una gráfica en donde salga el total de no2 de una estación a lo largo del tiempo indeterminado      
    elif estacion and not since:
        historico_df_filtradof = historico_df[historico_df["fiwareid"] == estacion]   #filtramos según la estación recibida

        plt.figure(figsize=(10,5))  #tamaño del gráfico (rectangular)
        #Marker muestra un circulo y linestyle une los circulos en una linea contínua
        plt.plot(historico_df_filtradof["fecha_carg"], historico_df_filtradof["no2"], marker='o', linestyle='-')    #dibuja gráfica del tiempo en donde "eje x" es la fecha de carga y el "eje y" el no2 generado.

        #Creamos gráfica
        plt.title(f"NO2 histórico por la estación {estacion} ")
        plt.xlabel("Fecha")
        plt.ylabel("Nivel NO2")
        plt.tight_layout()

        #La guardamos
        os.makedirs(output_dir, exist_ok=True)
        plt.savefig(os.path.join(output_dir, "NO2_historico_estacion.png"))
        plt.close()
        print("Graficos historicos guardados")
    #Crea una gráfica de barras de la media del NO2 de cada estación determinado por el tiempo
    elif since and not estacion:
        since_date = pd.to_datetime(since, utc=True) #pasamos since a una variable de tipo datetime utc
        historico_df_filtrados = historico_df[historico_df["fecha_carg"] >= since_date] #filtrado de fechas
        plt.figure(figsize=(15,10))  #tamaño del gráfico (rectangular)
        #Creamos gráfica
        historico_df_filtrados = historico_df_filtrados.groupby("fiwareid")["no2"].mean()
        historico_df_filtrados.plot(kind="bar", title=f"Gráfico de barras: media de no2 por estación desde {since} hasta la última fecha de carga", color = "red")
        plt.xlabel("Estacion")
        plt.ylabel("Media de NO2") 

        #La guardamos
        os.makedirs(output_dir, exist_ok=True)
        plt.savefig(os.path.join(output_dir, "NO2_historico_since.png"))
        plt.close()
        print("Graficos historicos guardados")
    else:
        plt.figure(figsize=(15,10))  #tamaño del gráfico (rectangular)
        #Creamos gráfica
        historico_df = historico_df.groupby("fiwareid")["no2"].mean()
        historico_df.plot(kind="bar", title=f"Gráfico de barras: media de no2 de todas las estaciones", color = "red")
        plt.xlabel("Estaciones")
        plt.ylabel("Media de NO2") 

        #La guardamos
        os.makedirs(output_dir, exist_ok=True)
        plt.savefig(os.path.join(output_dir, "NO2_historico.png"))
        plt.close()
        print("Graficos historicos guardados")
   

def main():
    args = parse_args()

    #Obtener datos de la API
    url = "https://valencia.opendatasoft.com/api/explore/v2.1/catalog/datasets/estacions-contaminacio-atmosferiques-estaciones-contaminacion-atmosfericas/records"
    try:
        response = requests.get(url, timeout=30)
        response.raise_for_status()
        data = response.json()  # Diccionario con más de una llave, solo queremos la llave 'results'
        print("Conexión exitosa a la API")
    except:
        print("Error en la conexión")
        return

    #Conexión a POSTGRES
    churro = 'postgresql://postgres:mysecretpassword@localhost:5432/postgres' #Puente de conexión que nos permite leer y escribir tablas en Postgres
    engine = create_engine(churro)


    save_raw_csv(data, engine)
    
    if args.modo == "historico":
        generate_historico_report(args.since, args.estacion)
    elif args.modo == "actual":
        generate_actual_report()

if __name__ == "__main__":
    main()

